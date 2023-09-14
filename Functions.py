import os
import pandas as pd
import numpy as np
import warnings
import pyreadstat
import statsmodels
import statsmodels.stats.weightstats
from scipy.stats import chi2_contingency
from itertools import combinations, product
from tqdm import tqdm
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

warnings.filterwarnings("ignore")


def analysis(file):
    assert file.lower().endswith(".sav"), "File must be an SPSS .SAV file."

    values, codebook = pyreadstat.read_sav(file, apply_value_formats=False)
    labels = pyreadstat.read_sav(file, apply_value_formats=True)[0]

    scale_variables = [
        key for key, measure in codebook.variable_measure.items() if measure == "scale"
    ]
    weights = ["Unweighted"] + [
        var for var in scale_variables if "weight" in var.lower()
    ]

    sample_comparisons = [
        comb
        for comb in list(
            combinations(
                list(
                    product(values["Group"].unique(), values["Time"].unique(), weights)
                ),
                2,
            )
        )
        if not (comb[0][0] == comb[1][0] and comb[0][2] != comb[1][2])
    ]

    for combination in tqdm(
        sample_comparisons, position=0, desc="Comparing Weighted Sampling", leave=True
    ):

        class sample:
            one = subsample(
                combination[0][0], combination[0][1], combination[0][2], values, labels
            )
            two = subsample(
                combination[1][0], combination[1][1], combination[1][2], values, labels
            )
            name = comparison_name(one, two)
            metadata = codebook
            paired = one.group + one.weight == two.group + two.weight

        sample.one.values.set_index(sample.one.values["IDs"], inplace=True)
        sample.two.values.set_index(sample.two.values["IDs"], inplace=True)
        sample.one.labels.set_index(sample.one.values["IDs"], inplace=True)
        sample.two.labels.set_index(sample.two.values["IDs"], inplace=True)

        analysis_tables(sample, "Nominal")
        analysis_tables(sample, "Ordinal")

    print("Analysis complete.")


def analysis_tables(sample, type):
    sample.crosstabs = pd.DataFrame()
    sample.summaries = pd.DataFrame()

    sample.metadata.variable_measure.pop("Group", None)
    sample.metadata.variable_measure.pop("Time", None)
    nominal_variables = [
        key
        for key, measure in sample.metadata.variable_measure.items()
        if measure == "nominal"
    ]

    if type == "Nominal":
        ordinal_variables = [1]
    else:
        ordinal_variables = list(
            reversed(
                [
                    key
                    for key, measure in sample.metadata.variable_measure.items()
                    if measure == "ordinal"
                ]
            )
        )

    for nominal_variable in tqdm(
        nominal_variables,
        position=1,
        desc="Comparing Nominal Variables",
        leave=False,
    ):
        for ordinal_variable in tqdm(
            ordinal_variables,
            position=2,
            desc="Comparing Ordinal Variables",
            leave=False,
            total=len(ordinal_variables),
        ):
            if type == "Nominal":
                sample.crosstab = nominal_crosstab(sample, nominal_variable)
                sample.summary = nominal_summary(sample, nominal_variable)
            if type == "Ordinal":
                sample.crosstab = ordinal_crosstab(
                    sample, nominal_variable, ordinal_variable
                )
                sample.summary = ordinal_summary(sample, ordinal_variable)

            sample.crosstabs = combine_crosstabs(sample.crosstab, sample.crosstabs)
            sample.summaries = pd.concat([sample.summaries, sample.summary], axis=0)

        if type == "Ordinal":
            second_header = [""] * len(sample.crosstabs.columns)
            second_header[2] = sample.one.name
            second_header[
                int((len(sample.crosstabs.columns) - 2) / 3) + 2
            ] = sample.two.name
            second_header[
                int((len(sample.crosstabs.columns) - 2) / 3) * 2 + 2
            ] = "Difference"
            multi_col = pd.MultiIndex.from_tuples(
                [
                    (second_header[i], col)
                    for i, col in enumerate(sample.crosstabs.columns)
                ]
            )
            sample.crosstabs.columns = multi_col

            write_xlsx(
                sample,
                document_title(
                    sample,
                    type,
                    sample.metadata.column_labels[
                        sample.metadata.column_names.index(nominal_variable)
                    ],
                ),
            )
            write_docx(
                sample,
                document_title(
                    sample,
                    type,
                    sample.metadata.column_labels[
                        sample.metadata.column_names.index(nominal_variable)
                    ],
                ),
                variable=sample.metadata.column_labels[
                    sample.metadata.column_names.index(nominal_variable)
                ],
            )
            sample.crosstabs = pd.DataFrame()
            sample.summaries = pd.DataFrame()

    if type == "Nominal":
        write_xlsx(sample, document_title(sample, type, nominal_variable))
        write_docx(sample, document_title(sample, type, nominal_variable))


def nominal_crosstab(sample, nominal_variable):
    sample.one.crosstab = create_crosstab(
        type="Nominal",
        data=sample.one.labels,
        index=nominal_variable,
        columns="Total",
        weight=sample.one.weight,
    )

    sample.two.crosstab = create_crosstab(
        type="Nominal",
        data=sample.two.labels,
        index=nominal_variable,
        columns="Total",
        weight=sample.two.weight,
    )

    sample.crosstab = pd.concat([sample.one.crosstab, sample.two.crosstab], axis=1)

    sample.crosstab = sample.crosstab.reset_index()
    sample.crosstab.insert(0, "Variable", np.nan)
    sample.crosstab.columns = [
        "Variable",
        "Values",
        add_sample_size(sample.one.name, sample.one.values[nominal_variable]),
        add_sample_size(sample.two.name, sample.two.values[nominal_variable]),
    ]

    sample.crosstab.loc[0, "Variable"] = test_chi(
        variable=sample.metadata.column_labels[
            sample.metadata.column_names.index(nominal_variable)
        ],
        observed=pd.crosstab(
            index=sample.one.labels[nominal_variable],
            columns="Total",
            values=sample.one.labels[sample.one.weight],
            aggfunc="sum",
        ),
        expected=pd.crosstab(
            index=sample.two.labels[nominal_variable],
            columns="Total",
            values=sample.two.labels[sample.one.weight],
            aggfunc="sum",
        ),
    )

    return sample.crosstab


def ordinal_crosstab(sample, nominal_variable, ordinal_variable):
    labels = sample.metadata.variable_value_labels[ordinal_variable].values()
    labels_ordered = []
    [labels_ordered.append(value) for value in labels if value not in labels_ordered]
    labels_ordered = labels_ordered + ["DK/NA"]

    sample.one.crosstab = create_crosstab(
        type="Ordinal",
        data=sample.one.labels,
        index=ordinal_variable,
        columns=nominal_variable,
        weight=sample.one.weight,
        labels=labels_ordered,
    )

    sample.two.crosstab = create_crosstab(
        type="Ordinal",
        data=sample.two.labels,
        index=ordinal_variable,
        columns=nominal_variable,
        weight=sample.two.weight,
        labels=labels_ordered,
    )

    sample.crosstab = pd.concat(
        [
            sample.one.crosstab,
            sample.two.crosstab,
            sample.two.crosstab - sample.one.crosstab,
        ],
        axis=1,
    )
    sample.crosstab = (
        sample.crosstab.round(1).apply(lambda x: x).applymap(lambda x: f"{x}%")
    )

    sample.crosstab = add_crosstab_tests(sample, nominal_variable, ordinal_variable)

    sample.crosstab = sample.crosstab.reset_index()
    sample.crosstab.insert(0, "Variable", np.nan)
    crosstab_header = sample.crosstab.columns.tolist()
    crosstab_header[0] = "Variable"
    crosstab_header[1] = "Label, Values"
    sample.crosstab.columns = crosstab_header
    sample.crosstab.iloc[0, 0] = ordinal_variable
    sample.crosstab.iloc[0, 1] = sample.metadata.column_labels[
        sample.metadata.column_names.index(ordinal_variable)
    ]

    return sample.crosstab


def ordinal_summary(sample, ordinal_variable):
    statement = ""
    sample.crosstab.set_index("Label, Values", inplace=True)

    label = sample.metadata.column_labels[
        sample.metadata.column_names.index(ordinal_variable)
    ]

    length = int((len(sample.crosstab.columns) - 1) / 3)

    for value in sample.crosstab.columns[1 : 1 + length]:
        index_one = list(sample.crosstab.columns).index(value)
        index_two = index_one + length
        index_dif = index_two + length

        value2 = sample.crosstab.columns[index_two]

        if type(sample.crosstab.iloc[0, index_dif]) == str:
            likeness, plurality1, plurality2 = plurality_comparison(
                sample.crosstab.iloc[:-1, index_one],
                sample.crosstab.iloc[:-1, index_two],
            )

            difference = mean_comparison(sample, sample.crosstab.iloc[0, index_dif])

            if sample.paired:
                if value.split(" (")[0] == "All":
                    statement += f"{sample.one.group} ({value.split(' (')[1]} responded to the statement, \"{label}\". The"
                else:
                    statement += f" Among those who selected \"{value.split(' (')[0]}\" ({value.split(' (')[1]}, the"
                statement += f" average response {difference} between {sample.one.time} and {sample.two.time}. At {sample.two.time}, there was {plurality2} among this group, {likeness} at {sample.one.time} {plurality1}"
                statement += f"."
            else:
                if value.split(" (")[0] == "All":
                    statement += f'{sample.one.name} and {sample.two.name} responded to the statement, "{label}". There'
                else:
                    statement += (
                        f" Among those who selected \"{value.split(' (')[0]}\", there"
                    )
                statement += f" {difference} in the average response between {sample.one.name} ({value.split(' (')[1]} and {sample.two.name} ({value2.split(' (')[1]}. Among {sample.two.group}, there was {plurality2}, {likeness} {sample.one.group} {plurality1}"
                statement += f"."
    if len(statement) == 0:
        statement = "Insufficient data to generate prose summary."
    return pd.DataFrame([ordinal_variable, statement]).T


def nominal_summary(sample, nominal_variable):
    statement = ""
    return pd.DataFrame([nominal_variable, statement]).T


def write_xlsx(sample, name):
    name += ".xlsx"
    title = sample.name

    os.makedirs(f"Outputs/{title}", exist_ok=True)

    sheet_name = title
    if len(sheet_name) > 31:
        sheet_name = sheet_name[:28] + "..."

    sample.crosstabs.index = [pd.NA] * len(sample.crosstabs)

    sample.crosstabs.to_excel(
        f"Outputs/{title}/Tables - {name}",
        sheet_name=sheet_name,
        index=True,
        header=True,
    )


def write_docx(sample, name, variable=None):
    name += ".docx"
    title = sample.name

    os.makedirs(f"Outputs/{title}", exist_ok=True)

    document = Document()

    section = document.sections[0]
    section.page_width = Inches(22)
    section.page_height = Inches(22)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    if variable == None:
        header = title
    else:
        header = f"{title} by {variable}"

    header = document.add_heading(header, 0)
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    sample.crosstabs.fillna("", inplace=True)

    document_summaries = document
    rows, cols = sample.summaries.shape
    table_summaries = document_summaries.add_table(rows=rows + 2, cols=cols)
    table_summaries.style = "Medium List 2"

    if 1 == 0:
        rows, cols = sample.crosstabs.shape
        table = document.add_table(rows=rows + 2, cols=cols)
        table.style = "Medium List 2"

        for i, column in enumerate(sample.crosstabs.columns):
            cell = table.cell(0, i)
            for p in cell.paragraphs:
                p.clear()
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if isinstance(column, tuple):
                p1 = cell.add_paragraph(str(column[0]))
                set_font(p1)
                p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                split = str(column[1]).find("(n = ")
                if split != -1:
                    parts = [str(column[1])[: split - 1], str(column[1])[split:]]
                    for part in parts:
                        p = cell.add_paragraph(part)
                        set_font(p)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    p2 = cell.add_paragraph(str(column[1]))
                    set_font(p2)
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                p = cell.add_paragraph(str(column))
                set_font(p)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i, row in enumerate(sample.crosstabs.iterrows()):
            data = row[1]
            for j, value in enumerate(data):
                cell = table.cell(i + 1, j)
                cell.text = str(value)
                set_font(cell)

                for paragraph in cell.paragraphs:
                    if j < 2:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    else:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                vertical_alignment(cell)

    for i, row in enumerate(sample.summaries.iterrows()):
        data = row[1]
        for j, value in enumerate(data):
            cell = table_summaries.cell(i + 1, j)
            cell.text = str(value)
            set_font(cell)

            for paragraph in cell.paragraphs:
                if j < 2:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            vertical_alignment(cell)

    document.save(f"Outputs/{title}/Tables - {name}")
    document_summaries.save(f"Outputs/{title}/Report - {name}")


def create_crosstab(type, data, index, columns, weight, labels=None):
    if type == "Nominal":
        margins = False
        dropna = True
        normalize = False
        index_data = data[index]
    else:
        margins = True
        dropna = False
        normalize = "columns"
        index_data = data[index].cat.add_categories(["DK/NA"]).fillna("DK/NA")

    absolute_frequencies = pd.crosstab(
        index=index_data,
        columns=data[columns],
        values=data[weight],
        aggfunc="sum",
        margins=margins,
        dropna=dropna,
        normalize=normalize,
    )

    if type == "Nominal":
        combined_frequencies = (
            (absolute_frequencies / absolute_frequencies.sum().sum() * 100)
            .round(1)
            .astype(str)
            .applymap(lambda x: f"({x}%)")
            + " "
            + absolute_frequencies.round().astype(int).astype(str)
        )

        return combined_frequencies.iloc[:, ::-1].replace("nan", "0")
    else:
        return 100 * absolute_frequencies.iloc[:, ::-1].fillna(0).reindex(labels)


def add_crosstab_tests(sample, nominal_variable, ordinal_variable):
    sample.means = pd.DataFrame(
        [[pd.NA] * len(sample.crosstab.columns)], columns=sample.crosstab.columns
    )

    for filter in sample.one.crosstab.columns:
        mean1, mean2, mean_difference = test_t(
            sample, filter, nominal_variable, ordinal_variable
        )

        crosstab_index = list(sample.one.crosstab.columns).index(filter)

        sample.means.iloc[
            0, crosstab_index + 0 * len(sample.one.crosstab.columns)
        ] = mean1
        sample.means.iloc[
            0, crosstab_index + 1 * len(sample.one.crosstab.columns)
        ] = mean2
        sample.means.iloc[
            0, crosstab_index + 2 * len(sample.one.crosstab.columns)
        ] = mean_difference

        crosstab_header = sample.crosstab.columns.tolist()
        crosstab_header[
            crosstab_index + 0 * len(sample.one.crosstab.columns)
        ] = add_sample_size(
            filter,
            sample.one.values
            if filter == "All"
            else sample.one.values[sample.one.labels[nominal_variable] == filter],
        )
        crosstab_header[
            crosstab_index + 1 * len(sample.one.crosstab.columns)
        ] = add_sample_size(
            filter,
            sample.two.values
            if filter == "All"
            else sample.two.values[sample.two.labels[nominal_variable] == filter],
        )
        sample.crosstab.columns = sample.means.columns = pd.Index(crosstab_header)

    return combine_crosstabs(sample.means, sample.crosstab)


def combine_crosstabs(crosstab1, crosstab2):
    if len(crosstab1.columns) != len(crosstab2.columns):
        crosstab2 = pd.DataFrame(columns=crosstab1.columns)

    crosstab1.columns = crosstab2.columns = pd.MultiIndex.from_product(
        [["Level1"], crosstab1.columns]
    )

    crosstabs = pd.concat([crosstab1, crosstab2])

    crosstabs.columns = crosstabs.columns.get_level_values(1)

    return crosstabs


def test_chi(variable, observed, expected):
    observed_expected = np.column_stack((observed, expected))

    observed_expected = observed_expected[
        ~np.apply_along_axis(lambda y: np.all(y == 0), 1, observed_expected)
    ]

    _, P, _, _ = chi2_contingency(observed_expected)

    if not np.isnan(P):
        if np.any(expected < 5):
            return f"{variable} (P = {P:.3f}) Warning: P-value may be incorrect because at least one expected value is less than 5."
        else:
            return f"{variable} (P = {P:.3f})"

    return variable


def test_t(sample, filter, nominal_variable, ordinal_variable):
    sample.one.filtered = sample.one.values
    sample.two.filtered = sample.two.values

    if not filter == "All":
        sample.one.filtered = sample.one.values[
            sample.one.labels[nominal_variable] == filter
        ]
        sample.two.filtered = sample.two.values[
            sample.two.labels[nominal_variable] == filter
        ]

    sample = full_entries(sample, ordinal_variable)

    if sample.paired:
        P = statsmodels.stats.weightstats.DescrStatsW(
            data=sample.two.ordinal_filtered - sample.one.ordinal_filtered,
            weights=sample.one.weights_filtered,
        ).ttest_mean(0)[1]

    else:
        P = statsmodels.stats.weightstats.ttest_ind(
            x1=sample.one.ordinal_filtered,
            x2=sample.two.ordinal_filtered,
            alternative="two-sided",
            usevar="unequal",
            weights=(sample.one.weights_filtered, sample.two.weights_filtered),
        )[1]

    if (
        not sum(sample.one.weights_filtered) == 0
        and not sum(sample.two.weights_filtered) == 0
    ):
        mean1 = np.average(
            sample.one.ordinal_filtered, weights=sample.one.weights_filtered
        )
        mean2 = np.average(
            sample.two.ordinal_filtered, weights=sample.two.weights_filtered
        )
        mean_difference = "{:.3f}".format(mean2 - mean1)
        mean1 = "{:.3f}".format(mean1)
        mean2 = "{:.3f}".format(mean2)
    else:
        mean1 = mean2 = mean_difference = pd.NA

    if not np.isnan(P):
        mean_difference = f"{mean_difference} (P = {P:.3f})"

    return mean1, mean2, mean_difference


def full_entries(sample, ordinal_variable):
    if sample.paired:
        complete_cases = pd.concat(
            (
                sample.one.filtered[ordinal_variable],
                sample.one.filtered[sample.one.weight],
                sample.two.filtered[ordinal_variable],
                sample.two.filtered[sample.two.weight],
            ),
            axis=1,
        ).dropna()
        sample.one.ordinal_filtered = complete_cases.groupby(level=0, axis=1).nth(0)[
            ordinal_variable
        ]
        sample.one.weights_filtered = complete_cases.groupby(level=0, axis=1).nth(0)[
            sample.one.weight
        ]
        sample.two.ordinal_filtered = complete_cases.groupby(level=0, axis=1).nth(1)[
            ordinal_variable
        ]
        sample.two.weights_filtered = complete_cases.groupby(level=0, axis=1).nth(1)[
            sample.two.weight
        ]

    else:
        complete_cases = pd.concat(
            (
                sample.one.filtered[ordinal_variable],
                sample.one.filtered[sample.one.weight],
            ),
            axis=1,
        ).dropna()
        sample.one.ordinal_filtered = complete_cases[ordinal_variable]
        sample.one.weights_filtered = complete_cases[sample.one.weight]

        complete_cases = pd.concat(
            (
                sample.two.filtered[ordinal_variable],
                sample.two.filtered[sample.two.weight],
            ),
            axis=1,
        ).dropna()
        sample.two.ordinal_filtered = complete_cases[ordinal_variable]
        sample.two.weights_filtered = complete_cases[sample.two.weight]

    return sample


def document_title(sample, type, nominal_variable):
    if type == "Nominal":
        return " - ".join([f"{type} Variables", sample.name])
    else:
        return " - ".join([f"{type} Variables", sample.name, nominal_variable])


def comparison_name(sample1, sample2):
    if sample1.weight == sample2.weight:
        if sample1.group == sample2.group:
            return f"{sample1.group} at {sample1.time} v. {sample2.time} ({sample1.weight})"
        elif sample1.time == sample2.time:
            return f"{sample1.group} v. {sample2.group} at {sample1.time} ({sample1.weight})"
        else:
            return f"{sample1.group} at {sample1.time} v. {sample2.group} at {sample2.time} ({sample1.weight})"
    else:
        if sample1.group == sample2.group:
            return f"{sample1.group} at {sample1.time} ({sample1.weight}) v. {sample2.time} ({sample2.weight})"
        elif sample1.time == sample2.time:
            return f"{sample1.group} ({sample1.weight}) v. {sample2.group} ({sample2.weight}) at {sample1.time}"
        else:
            return f"{sample1.group} at {sample1.time} ({sample1.weight}) v. {sample2.group} at {sample2.time} ({sample2.weight})"


def add_sample_size(variable, sample):
    return f"{variable} (n = {len(sample)})"


def set_font(item):
    if hasattr(item, "paragraphs"):  # Check if it's a cell
        for paragraph in item.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
    else:
        for run in item.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def vertical_alignment(cell, align="bottom"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def mean_comparison(sample, difference):
    numeric = float(difference.split(" (")[0])
    if "P =" in difference:
        P = float(difference.split("P = ")[1][:-1])
    else:
        P = 1

    if numeric < 0:
        change = "decreased"
    elif numeric > 0:
        change = "increased"
    else:
        change = "did not change"

    if sample.paired:
        if P < 0.05:
            return f"{change} significantly by {difference}"
        else:
            return f"did not change significantly (P = {P})"
    else:
        if P < 0.05:
            return f"was a significant difference by {difference}"
        else:
            return f"was not a significant (P = {P}) difference"


def plurality(percentages):
    value = percentages.str.rstrip("%").astype(float).idxmax()
    numeric = float(percentages[value].rstrip("%"))

    if numeric / 100 < 0.5:
        kind = "plurality"
    elif numeric / 100 > 2 / 3:
        kind = "supermajority"
    else:
        kind = "majority"

    return value, numeric, kind


def plurality_comparison(percentages1, percentages2):
    value1, numeric1, kind1 = plurality(percentages1)
    value2, numeric2, kind2 = plurality(percentages2)

    if kind1 == kind2:
        likeness = "like"
        plurality1 = f"({numeric1}%)"
    else:
        likeness = "unlike"
        plurality1 = f'a {kind1} response of "{value1}" ({numeric1}%)'

    plurality2 = f'a {kind2} response of "{value2}" ({numeric2}%)'

    return likeness, plurality1, plurality2


class subsample:
    def __init__(self, group, time, weight, values, labels):
        self.group = group
        self.time = time
        self.weight = weight
        self.name = f"{group} at {time}"
        self.values = (
            values[(values["Group"] == group) & (values["Time"] == time)]
            .assign(Total="Total")
            .assign(Unweighted=1)
        )
        self.labels = (
            labels[(values["Group"] == group) & (values["Time"] == time)]
            .assign(Total="Total")
            .assign(Unweighted=1)
        )


analysis("Dataset copy.sav")
